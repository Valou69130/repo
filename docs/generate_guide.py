"""
CollateralOS — User & Technical Guide generator.
Run: python3 docs/generate_guide.py
Output: docs/CollateralOS_Guide.docx
Update this file whenever a new feature ships.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ────────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()
    return table

def hr(doc):
    doc.add_paragraph("─" * 80)

# ── document ───────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ── Cover ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_heading("CollateralOS", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("User & Technical Guide")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(16)
sub.runs[0].bold = True

ver = doc.add_paragraph(f"Version 2026.04  ·  {datetime.date.today().strftime('%-d %B %Y')}  ·  Romania Pilot")
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_page_break()

# ── 1. Introduction ────────────────────────────────────────────────────────────
add_heading(doc, "1. Introduction")
add_para(doc,
    "CollateralOS is a full-stack collateral management platform built for treasury and "
    "operations teams at financial institutions. It covers the complete lifecycle of "
    "collateral: inventory, repo transactions, margin monitoring, margin calls, dispute "
    "resolution, four-eyes approvals, regulatory reporting, and AI-assisted analysis.",
)
add_para(doc,
    "This document serves two audiences:",
)
add_bullet(doc, "End users (treasury managers, collateral managers, operations analysts, "
                "risk reviewers, credit approvers) — explains what each module does and "
                "how to use it.")
add_bullet(doc, "Technical readers (developers, integration teams) — explains the system "
                "architecture, API design, data model, and extension points.")

doc.add_paragraph()
add_heading(doc, "1.1  Production URL", level=2)
add_para(doc, "https://www.collateralos.app")

add_heading(doc, "1.2  Tech Stack", level=2)
add_table(doc,
    ["Layer", "Technology"],
    [
        ["Frontend",       "React 18, Vite, Tailwind CSS, shadcn/ui"],
        ["Backend API",    "Node.js, Express 4, better-sqlite3"],
        ["Database",       "SQLite (file-based, WAL mode, hash-chained audit)"],
        ["Auth",           "JWT (HTTP-only cookies, access + refresh tokens)"],
        ["AI",             "Anthropic Claude API (claude-3-5-sonnet)"],
        ["Deployment",     "Vercel (serverless, /tmp SQLite per instance)"],
    ],
)

doc.add_page_break()

# ── 2. User Roles ──────────────────────────────────────────────────────────────
add_heading(doc, "2. User Roles & Permissions")
add_para(doc,
    "CollateralOS uses role-based access control. Each demo user maps to a specific "
    "function within the institution."
)
add_table(doc,
    ["Role", "Email (demo)", "Typical responsibilities"],
    [
        ["Treasury Manager",    "treasury@banca-demo.ro",    "Overall portfolio view, EoD lock, business case"],
        ["Collateral Manager",  "collateral@banca-demo.ro",  "Inventory, repo management, substitutions"],
        ["Operations Analyst",  "operations@banca-demo.ro",  "Settlement, SFTR, compliance, integration"],
        ["Risk Reviewer",       "risk@banca-demo.ro",        "Margin monitoring, parameters, audit review"],
        ["Credit Approver",     "approver@banca-demo.ro",    "Four-eyes approval on large margin calls"],
    ],
)
add_para(doc, "All demo accounts use the password: demo1234", italic=True,
         color=(0x64, 0x74, 0x8b))

doc.add_page_break()

# ── 3. Module Guide ────────────────────────────────────────────────────────────
add_heading(doc, "3. Module Guide")

# Dashboard
add_heading(doc, "3.1  Dashboard", level=2)
add_para(doc,
    "The landing page after login. Shows KPI cards (total collateral, active repos, "
    "margin at risk, pending notifications), a Suggested Calls panel that surfaces "
    "agreements with a collateral deficit, and a Pending Approvals widget for the "
    "four-eyes inbox."
)
add_para(doc, "Key interactions:", bold=True)
add_bullet(doc, "Click a suggested call row to open the Agreement detail.")
add_bullet(doc, "Click a pending approval to jump to the Approvals page.")
add_bullet(doc, "Use the Top Bar to switch roles (demo only) or reset the demo data.")

# Collateral Inventory
add_heading(doc, "3.2  Collateral Inventory", level=2)
add_para(doc,
    "Displays all assets held in custody. Each asset record carries: ISIN, issuer, type, "
    "currency, quantity, market value, haircut %, eligibility classification, custody "
    "location, and status (Available / Pledged / Pending)."
)
add_para(doc, "Key interactions:", bold=True)
add_bullet(doc, "Edit any asset inline to update market value, haircut, or eligibility.")
add_bullet(doc, "Import assets via CSV upload (download the template from the import dialog).")
add_bullet(doc, "Search and filter by type, currency, custody, or status.")

# Repo Transactions
add_heading(doc, "3.3  Repo Transactions", level=2)
add_para(doc,
    "Manages the lifecycle of repo / reverse-repo agreements: Open → Active → Closed. "
    "Each repo links to a collateral asset, records the haircut-adjusted posted amount, "
    "and tracks maturity."
)
add_para(doc, "Key interactions:", bold=True)
add_bullet(doc, "Create a new repo from the 'New Repo' button.")
add_bullet(doc, "Open a repo detail to top-up collateral, propose a substitution, or roll over.")
add_bullet(doc, "Substitution proposals require a second user to approve (four-eyes).")

# Margin Monitor
add_heading(doc, "3.4  Margin Monitor", level=2)
add_para(doc,
    "Shows real-time margin exposure across all active repos. Flags repos where the "
    "mark-to-market move exceeds the threshold set in Parameters & Rules. "
    "AI-assisted analysis is available per repo."
)

# Agreements & Margin Calls
add_heading(doc, "3.5  Agreements & Margin Calls", level=2)
add_para(doc,
    "Agreements are GMRA / CSA master agreements between the institution and a "
    "counterparty. Each agreement holds a threshold, minimum transfer amount (MTA), "
    "and rounding rules. Margin calls are issued against an agreement when the net "
    "exposure exceeds the threshold."
)
add_para(doc, "Margin call lifecycle states:", bold=True)
add_table(doc,
    ["State", "Description"],
    [
        ["draft",             "Created but not yet sent to counterparty"],
        ["issued",            "Formally issued; counterparty has been notified"],
        ["pending_four_eyes", "Accepted above threshold; awaiting second-user approval"],
        ["agreed",            "Both parties agreed on the call amount"],
        ["disputed",          "Counterparty has raised a dispute"],
        ["delivered",         "Collateral has been transferred"],
        ["settled",           "Settlement confirmed by both sides"],
        ["cancelled",         "Call withdrawn"],
    ],
)
add_para(doc, "Key interactions:", bold=True)
add_bullet(doc, "From an Agreement detail, issue a new margin call with a proposed amount.")
add_bullet(doc, "Counterparty response (accept / dispute) moves the call through the lifecycle.")
add_bullet(doc, "Calls above the four-eyes threshold route to the Approvals inbox automatically.")

# Disputes
add_heading(doc, "3.6  Disputes", level=2)
add_para(doc,
    "When a counterparty disagrees with a margin call amount, they open a dispute. "
    "The dispute workflow supports propose / agree / withdraw / escalate actions. "
    "All state transitions are hash-chained in the audit log."
)
add_para(doc, "Dispute states:", bold=True)
add_table(doc,
    ["State", "Description"],
    [
        ["open",      "Counterparty has raised an objection"],
        ["proposed",  "Institution has put forward a settlement amount"],
        ["agreed",    "Both parties accepted the proposed amount"],
        ["withdrawn", "Dispute withdrawn; original call stands"],
        ["escalated", "Sent to senior management / legal"],
    ],
)

# Four-Eyes Approvals
add_heading(doc, "3.7  Four-Eyes Approvals", level=2)
add_para(doc,
    "Any margin call accepted above the configured threshold (default: 500,000 RON) "
    "is placed in a pending_four_eyes state and routed to the Approvals inbox. "
    "A second authorised user (Credit Approver or Treasury Manager) must grant or "
    "reject the approval before the call progresses."
)
add_bullet(doc, "The original submitter cannot approve their own call.")
add_bullet(doc, "Approval and rejection are both recorded in the audit log with reason.")

# Operations
add_heading(doc, "3.8  Settlement / Operations", level=2)
add_para(doc,
    "Tracks settlement status of repos. Operations analysts confirm delivery and "
    "settlement. The ops scanner runs background checks for overdue deliveries and "
    "flags exceptions."
)

# Audit Trail & Audit Export
add_heading(doc, "3.9  Audit Trail & Audit Export", level=2)
add_para(doc,
    "Every state-changing action is written to the audit_events table with a SHA-256 "
    "hash chained to the previous entry. This makes the log tamper-evident: any "
    "modification to a historical record breaks the chain."
)
add_para(doc, "Key interactions:", bold=True)
add_bullet(doc, "Audit Trail page: search and filter all events by user, role, or action type.")
add_bullet(doc, "Audit Export page: download a signed CSV of events for a date range; "
                "verify the chain integrity with the 'Verify Chain' button.")

# AI Features
add_heading(doc, "3.10  AI Features", level=2)
add_para(doc,
    "Requires ANTHROPIC_API_KEY to be set in the API environment. When enabled, "
    "the following AI-powered tools are available:"
)
add_table(doc,
    ["Feature", "Where", "Description"],
    [
        ["Explain Deficit",      "Margin Monitor / Repo Detail", "Natural-language explanation of why a margin deficit exists"],
        ["Analyse Portfolio",    "Margin Monitor",               "Portfolio-level concentration and correlation analysis"],
        ["Correlate Exceptions", "Operations",                   "Groups related margin exceptions by likely root cause"],
        ["AI Assess Call",       "Margin Call Detail",           "Rates a margin call for urgency and recommends action"],
        ["AI Chat",              "Top Bar (Sparkles icon)",      "Free-form assistant with full portfolio context"],
    ],
)

# Parameters & Rules
add_heading(doc, "3.11  Parameters & Rules", level=2)
add_para(doc,
    "Configures the eligibility schedules and haircut schedules that drive collateral "
    "valuation and margin calculations. Changes are applied immediately to all "
    "subsequent margin calculations."
)

doc.add_page_break()

# ── 4. Technical Architecture ──────────────────────────────────────────────────
add_heading(doc, "4. Technical Architecture")

add_heading(doc, "4.1  Repository Layout", level=2)
add_table(doc,
    ["Path", "Description"],
    [
        ["collateral-app/",         "React frontend (Vite)"],
        ["collateral-app/src/pages/",  "One file per top-level view"],
        ["collateral-app/src/components/", "Shared UI components"],
        ["collateral-app/src/lib/api.js",  "HTTP client (all /api/* calls)"],
        ["collateral-app/src/domain/",     "Permissions, formatters, store"],
        ["collateral-api/",         "Express backend"],
        ["collateral-api/src/routes/",  "One file per resource group"],
        ["collateral-api/src/db/",      "SQLite schema, seed data, appendEvent"],
        ["collateral-api/src/middleware/", "JWT auth, role guard"],
        ["collateral-api/src/ai/",    "Claude integration, proactive scheduler"],
        ["api/[...route].js",        "Vercel catch-all — proxies all /api/* to Express"],
        ["vercel.json",              "Build, output, and routing config"],
        ["docs/",                    "This guide and Superpowers specs/plans"],
    ],
)

add_heading(doc, "4.2  API Design", level=2)
add_para(doc,
    "All API routes are mounted under /api on the Vercel deployment. Locally the "
    "Express server runs on port 3001. Auth uses HTTP-only cookies (co_token for the "
    "2-hour access JWT, co_refresh for the 7-day refresh JWT). Every mutating endpoint "
    "requires requireAuth middleware."
)
add_table(doc,
    ["Route group", "Mount path", "Description"],
    [
        ["Auth",        "/api/session",       "Login, logout, /me, refresh, password change"],
        ["Assets",      "/api/assets",        "CRUD + CSV import"],
        ["Repos",       "/api/repos",         "Repo lifecycle + settlement"],
        ["Agreements",  "/api/agreements",    "GMRA/CSA master agreements"],
        ["Margin Calls","/api/margin-calls",  "Event-sourced call lifecycle"],
        ["Disputes",    "/api/disputes",      "Dispute workflow"],
        ["Approvals",   "/api/approvals",     "Four-eyes inbox"],
        ["Audit",       "/api/audit",         "Read audit log, verify chain"],
        ["Notifications","/api/notifications","CRUD + acknowledge"],
        ["Admin",       "/api/admin",         "Demo reset, CSV template"],
        ["AI",          "/api/ai",            "Explain, analyse, chat, scheduler status"],
    ],
)

add_heading(doc, "4.3  Event Sourcing (Margin Calls & Disputes)", level=2)
add_para(doc,
    "Margin calls and disputes are event-sourced. The current state is derived by "
    "replaying events from the margin_call_events table. Each event carries:"
)
add_bullet(doc, "call_id, event_type (issued / accepted / agreed / cancelled / …)")
add_bullet(doc, "actor_user_id, actor_role")
add_bullet(doc, "payload (JSON — amount, reason, etc.)")
add_bullet(doc, "prev_hash (SHA-256 of the previous event — tamper-evident chain)")
add_bullet(doc, "ts (UTC timestamp)")
add_para(doc,
    "The appendEvent() function in collateral-api/src/db/appendEvent.js handles "
    "transition validation, hash chaining, and four-eyes routing in a single SQLite "
    "transaction.",
)

add_heading(doc, "4.4  Authentication Flow", level=2)
add_para(doc,
    "1. POST /api/session/login → verifies bcrypt hash, issues access + refresh JWT "
    "as HTTP-only SameSite=Strict cookies.\n"
    "2. Every subsequent request reads co_token from the cookie.\n"
    "3. On 401, the frontend transparently calls POST /api/session/refresh "
    "using co_refresh, obtains a new access token, and retries once.\n"
    "4. If refresh also fails, the user is logged out and redirected to login."
)

add_heading(doc, "4.5  Vercel Deployment Notes", level=2)
add_para(doc, "Important constraints for the serverless deployment:", bold=True)
add_bullet(doc, "SQLite lives in /tmp/collateral-demo.db per Lambda instance. "
                "Each cold start seeds fresh demo data. Data does NOT persist across restarts.")
add_bullet(doc, "better-sqlite3 is a native module — vercel.json runs "
                "'npm rebuild better-sqlite3' after install to compile for Linux.")
add_bullet(doc, "vercel.json routes /api/(.*) explicitly to the catch-all function "
                "because Vercel's default file-system routing does not forward nested "
                "paths to [...]route.js reliably.")
add_bullet(doc, "Vercel intercepts /api/auth/* and /api/account/* at the edge. "
                "Auth is therefore mounted at /api/session/*.")
add_bullet(doc, "ALLOWED_ORIGINS, JWT_SECRET, JWT_REFRESH_SECRET, and optionally "
                "ANTHROPIC_API_KEY must be set as Vercel environment variables.")

doc.add_page_break()

# ── 5. Environment Variables ───────────────────────────────────────────────────
add_heading(doc, "5. Environment Variables")
add_table(doc,
    ["Variable", "Required", "Description"],
    [
        ["JWT_SECRET",           "Yes", "Signs access tokens (min 32 chars)"],
        ["JWT_REFRESH_SECRET",   "Yes", "Signs refresh tokens (different from JWT_SECRET)"],
        ["ALLOWED_ORIGINS",      "Yes", "Comma-separated list of allowed CORS origins"],
        ["ANTHROPIC_API_KEY",    "No",  "Enables AI features; AI is disabled if absent"],
        ["COLLATERAL_DB_PATH",   "No",  "Override SQLite file path (default: /tmp/... on Vercel)"],
        ["NODE_ENV",             "No",  "Set to 'production' to enable combined HTTP logs"],
        ["PORT",                 "No",  "Local dev server port (default: 3001)"],
    ],
)

doc.add_page_break()

# ── 6. Demo Reset ──────────────────────────────────────────────────────────────
add_heading(doc, "6. Demo Reset")
add_para(doc,
    "The demo can be reset to its initial seeded state at any time. In the app, "
    "click the user menu in the top-right corner and select 'Reset Demo'. This calls "
    "POST /api/admin/reset, which re-runs the full seed: 3 agreements, 5 margin calls "
    "with complete event histories, demo assets, repos, and all 5 users."
)
add_para(doc,
    "Because Vercel SQLite is ephemeral (/tmp), the demo also auto-resets on each "
    "cold start of the serverless function.",
    italic=True, color=(0x64, 0x74, 0x8b),
)

doc.add_page_break()

# ── 7. Updating This Document ──────────────────────────────────────────────────
add_heading(doc, "7. Keeping This Document Up to Date")
add_para(doc,
    "This document is generated by docs/generate_guide.py. When a new feature ships:"
)
add_bullet(doc, "Add or update the relevant section in generate_guide.py.")
add_bullet(doc, "Run:  python3 docs/generate_guide.py")
add_bullet(doc, "Commit the updated docs/CollateralOS_Guide.docx alongside the feature.")

# ── footer ─────────────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph(
    f"CollateralOS  ·  Romania Pilot  ·  Generated {datetime.date.today().strftime('%-d %B %Y')}  ·  "
    "https://www.collateralos.app"
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
p.runs[0].font.size = Pt(9)

# ── save ────────────────────────────────────────────────────────────────────────
out = "docs/CollateralOS_Guide.docx"
doc.save(out)
print(f"Saved → {out}")
